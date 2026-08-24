"""Create the IPO service data governance and backend pipeline design document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "IPO_데이터_거버넌스_및_백엔드_파이프라인_설계서.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "NanumGothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "NanumGothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "IPO Prediction Service | Data Governance & Backend Pipeline"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(section.footer.paragraphs[0])


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "NanumGothic")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    header_row_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_row_properties.append(header_marker)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header_row_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_row_properties.append(header_marker)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_source(doc, text, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text + " ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run = p.add_run(url)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(46, 116, 181)


def build_document():
    doc = Document()
    configure_document(doc)
    add_title(doc, "IPO 데이터 거버넌스 및 백엔드 파이프라인 설계서", "공모주 상장 전 시초가·정규장 종가 수익률 예측 서비스 | 초기 운영 설계")

    add_table(doc, ["문서 항목", "내용"], [
        ("목적", "실제 DART·KRX 데이터 기반 예측 서비스를 출시 가능한 운영 시스템으로 설계한다."),
        ("적용 범위", "데이터 수집, 저장, 품질, 접근제어, 모델 학습·배포, 상장일 실측 확정, 감사와 장애 대응"),
        ("핵심 원칙", "상장 전 시점성 보장, 원문 보존, 최소권한, 재현성, 품질 우선, 출처·이용권한 확인"),
        ("문서 기준일", date.today().isoformat()),
    ], [1800, 7560])

    add_callout(doc, "결정 요약", "이 서비스는 ‘DART 공시로 상장 전에 알 수 있는 정보’를 피처로 삼고 ‘KRX 정규장 가격’으로 실제 결과를 확정한다. 따라서 원문, 정규화 데이터, 피처 스냅샷, 모델 버전, 예측값을 서로 연결해 보관해야 한다. 원시 데이터만 맞아도 시점이 틀리면 모델 성능은 신뢰할 수 없다.")

    doc.add_heading("1. 목적과 설계 범위", level=1)
    doc.add_paragraph("서비스의 목적은 공모주가 상장되기 전에 확정 공모가를 기준으로 상장일 시초가 수익률과 KRX 정규장 종가 수익률을 각각 예측하고, 상장일 21:00 KST에 실측·오차를 확정하는 것이다. 서비스는 투자 권유가 아니라 데이터 기반 예측 정보를 제공하며, 출처와 갱신 시점, 불확실성 구간을 함께 표시한다.")
    add_bullet(doc, "예측 타깃: (시초가 / 확정 공모가 - 1) x 100, (정규장 종가 / 확정 공모가 - 1) x 100")
    add_bullet(doc, "상장 전 피처: 증권신고서, 공모가 밴드·확정가, 수요예측, 기관 의무보유확약, 신주·구주, 최대주주 보호예수, 재무제표, 직전 시장·섹터 지표")
    add_bullet(doc, "상장 후 실측: KRX 상장일 시초가·정규장 종가. NXT 애프터마켓은 관측 데이터로 분리하고 모델 타깃과 혼합하지 않는다.")

    doc.add_heading("2. 인증정보를 시스템에 설정하는 방법", level=1)
    doc.add_paragraph("인증정보는 코드, Git, 문서, 로그에 직접 넣지 않는다. 현재 코드는 운영체제 환경변수에서 값을 읽으므로 개발 PC에서는 터미널 환경변수로 주입한다. 향후 .env 방식을 도입할 경우에는 전용 로더를 추가하고 Git에 포함하지 않는다. 운영 환경은 클라우드 Secret Manager 또는 배포 플랫폼의 Secret 기능을 사용한다. 애플리케이션은 시작할 때만 비밀값을 읽고 이후에는 값 자체를 출력하지 않는다.")
    add_table(doc, ["환경", "설정 방식", "필수 값", "확인 방법"], [
        ("개발 PC", "터미널 환경변수 (현재 코드)", "DART_API_KEY, KRX_ID, KRX_PASSWORD", "collect 명령이 키 누락 오류 없이 시작되는지 확인"),
        ("CI/CD", "암호화된 저장소 Secret", "동일", "값을 출력하지 않는 연결·단위 테스트"),
        ("운영", "Secret Manager + 서비스 계정", "동일 및 DB/서명 키", "헬스체크는 설정 여부만 보고"),
    ], [1200, 2700, 2800, 2660])
    doc.add_heading("2.1 개발 PC 예시", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("export DART_API_KEY=발급받은_40자리_키\nexport KRX_ID=개인_KRX_아이디\nexport KRX_PASSWORD=개인_KRX_비밀번호\npython pipeline.py --mode collect --start-year 2020 --end-year 2025 --phase phase2")
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    add_callout(doc, "운영 규칙", "DART 키는 OpenDART에서 발급받는다. KRX 계정·시장 데이터의 이용 범위와 재배포 가능 여부는 서비스 출시 전 KRX/코스콤 약관 및 계약으로 확정한다. 특히 앱에 원시 시세를 재배포하는 기능은 별도 법무·데이터 라이선스 검토 없이는 출시하지 않는다.")

    doc.add_heading("3. 통합 데이터 거버넌스 프레임워크", level=1)
    doc.add_paragraph("본 서비스는 중앙 기준과 데이터 도메인 책임을 함께 두는 하이브리드 모델을 채택한다. 제품 책임자가 정책과 출시 결정을 맡고, 데이터 책임자가 원천·품질·계보를, 모델 책임자가 재현성·성능·편향 및 배포 승인을 맡는다. SailPoint의 소유권·접근제어·보존·품질·스튜어드십 관점과 Snowflake의 책임성·메타데이터·감사 증거 관점을 서비스 규모에 맞게 적용했다.")
    add_table(doc, ["역할", "주요 책임", "승인 권한"], [
        ("서비스 오너", "서비스 범위, 데이터 이용 목적, 사용자 고지, 출시/중단", "정책·출시 최종 승인"),
        ("데이터 오너", "DART/KRX 원천 계약·정의·보존 기간·품질 기준", "소스 추가 및 데이터 계약 승인"),
        ("데이터 스튜어드", "필드 사전, 매칭 오류, 결측/중복 처리, 품질 이슈", "데이터 격리·복구 요청"),
        ("모델 오너", "피처 시점성, 학습·백테스트, 모델 카드, 승격 기준", "모델 후보 배포 요청"),
        ("플랫폼/보안 오너", "비밀관리, 권한, 암호화, 감사로그, 백업", "운영 접근권한 승인"),
    ], [1600, 4960, 2800])

    doc.add_heading("3.1 데이터 분류와 보존", level=2)
    add_table(doc, ["등급", "예시", "통제", "기본 보존"], [
        ("공개", "DART 공시 식별자, 공개된 기업 정보", "읽기 전용 원문 해시·출처 기록", "7년, 이용권한 재검토"),
        ("제한", "KRX 수집 데이터, 가공 시계열, 원문 파일", "서비스 계정만 수집, 사용자 직접 다운로드 차단", "계약·약관 및 내부 정책 기준"),
        ("기밀", "API 키, KRX ID·비밀번호, DB 비밀번호", "Secret Manager, 마스킹, 로그 금지, 90일 검토", "회전 즉시 이전 버전 폐기"),
        ("내부", "피처 스냅샷, 모델, 예측·오차, 감사 로그", "역할 기반 접근, 버전 고정", "모델·감사 목적상 7년 / 로그 2년"),
    ], [1000, 2900, 3300, 2160])

    doc.add_heading("3.2 데이터 카탈로그와 계보", level=2)
    doc.add_paragraph("모든 테이블과 파일에 데이터셋 ID, 소유자, 원천 URL·수집 시각, 수집 워커 버전, 스키마 버전, 민감도, 품질 상태, 보존 정책, 이용 제약을 기록한다. 한 예측 결과는 반드시 ‘어느 원문 공시의 어느 버전 → 어떤 정규화 규칙 → 어떤 피처 스냅샷 → 어떤 모델’로 추적 가능해야 한다.")
    add_bullet(doc, "원문: rcept_no, 문서 해시(SHA-256), 수집 시각, 원문 저장 위치")
    add_bullet(doc, "정규화: source_record_id, 변환 코드 버전, 파싱 성공 여부, 값의 근거 텍스트 위치")
    add_bullet(doc, "피처: as_of_at, feature_set_version, 데이터 절단 시각, 결측 사유")
    add_bullet(doc, "모델/예측: model_version, train_data_version, prediction_id, 입력 스냅샷 ID, 신뢰구간")

    doc.add_heading("4. 백엔드 아키텍처", level=1)
    doc.add_paragraph("초기 출시 구조는 모듈형 모놀리식 API와 비동기 워커로 시작하고, 데이터량·사용자 수가 늘면 수집/학습/서빙을 독립 서비스로 분리한다. 핵심은 서비스 수를 늘리는 것이 아니라 각 데이터 상태가 명확히 분리되는 것이다.")
    add_table(doc, ["계층", "구성", "책임"], [
        ("API 계층", "FastAPI, 인증, 사용자·관리자 API", "예측 조회, 데이터 최신성·면책 고지, 권한 검사"),
        ("작업 계층", "스케줄러 + 큐 워커", "DART/KRX 수집, 파싱, 정합, 21:00 실측, 재학습"),
        ("도메인 계층", "IPO, Feature, Model, Evaluation 서비스", "업무 규칙, 시점성 검증, 모델 승격"),
        ("저장 계층", "PostgreSQL, 객체 저장소, 모델 레지스트리", "정규화 데이터, 원문, 피처 스냅샷, 모델·감사 보존"),
        ("관측 계층", "로그, 메트릭, 알림", "수집 실패, 품질 저하, 예측 지연, 모델 드리프트 감시"),
    ], [1500, 3100, 4760])
    add_callout(doc, "권장 저장소 분리", "원문 ZIP/XML과 대용량 수집 결과는 객체 저장소에 불변으로 저장한다. PostgreSQL에는 조회·운영에 필요한 메타데이터와 정규화 레코드만 둔다. 피처 파일은 학습 재현성을 위해 스냅샷 ID와 함께 Parquet로 보존한다.")

    doc.add_heading("5. 백엔드 파이프라인 프로세스", level=1)
    add_table(doc, ["단계", "입력", "처리와 통제", "출력"], [
        ("1. 일정 탐지", "KRX 상장 기본정보, 내부 IPO 후보", "신규/변경 탐지, 중복 키 검사, 데이터 이용권한 상태 검사", "ipo_candidate"),
        ("2. 공시 수집", "DART 공시 목록·원문", "인증키 사용, 원문 해시, ZIP 보관, 호출 제한·재시도", "raw_disclosure"),
        ("3. 파싱·정규화", "원문·재무 API", "필드 단위 근거, 단위 통일, 파싱 실패는 결측으로 유지", "ipo_disclosure_normalized"),
        ("4. 정합", "DART + KRX", "corp_code·종목코드·상장일·법인명 검증, 모호한 매칭은 격리", "ipo_master"),
        ("5. 피처 스냅샷", "정규화 데이터·전일 시장지수", "as_of_at 이전 데이터만 허용, 누수 검사, 품질 게이트", "feature_snapshot"),
        ("6. 예측", "승격 모델 + 스냅샷", "모델 버전·입력 ID·신뢰구간 기록", "prediction"),
        ("7. 실측 확정", "KRX 시초가·정규장 종가", "상장일 21:00 KST 수집, 재조회·불일치 검토", "outcome/evaluation"),
        ("8. 재학습", "누적 성과", "10건 추가 또는 월 1회, 시계열 검증 우위일 때만 승격", "candidate_model/model_registry"),
    ], [1200, 1900, 3750, 2510])

    doc.add_heading("5.1 운영 스케줄과 상태 전이", level=2)
    add_number(doc, "매일 오전: 예정 상장·공시 변경을 탐지하고 후보 종목 상태를 갱신한다.")
    add_number(doc, "수요예측·공모가 확정 공시 탐지 시: 원문을 보관하고 피처 스냅샷을 새로 만든다.")
    add_number(doc, "상장 전: 필수 피처 품질 게이트를 통과한 스냅샷만 예측에 사용한다. 미통과 종목은 ‘데이터 불충분’ 상태로 표시한다.")
    add_number(doc, "상장일 21:00 KST: KRX 정규장 시초가·종가를 확정하고 예측 오차를 기록한다. NXT 애프터마켓은 별도 관측 필드로 보관한다.")
    add_number(doc, "매월 또는 신규 10개 실측 후: 후보 모델을 과거 시점 순서대로 검증하고 현행 모델보다 개선된 경우만 승격한다.")

    doc.add_heading("6. 데이터 품질과 시점성 통제", level=1)
    doc.add_paragraph("데이터 품질은 단순 결측률이 아니라 모델이 실제 상장 전에 이용할 수 있었는지를 판단하는 운영 기준이다. 실패한 레코드는 조용히 0으로 채우지 않고 사유 코드와 함께 격리한다.")
    add_table(doc, ["품질 차원", "규칙", "실패 시 조치", "KQI"], [
        ("정확성", "공모가·시초가·종가는 양수이며 원문/원천과 일치", "원천 재조회, 수동 검토", "원천 대조 일치율"),
        ("완전성", "필수 피처와 타깃의 확보 상태 기록", "예측 제외 또는 불확실성 경고", "필수 피처 충족률"),
        ("적시성", "feature.as_of_at < listing_date 및 공시 접수 시각", "스냅샷 폐기·재생성", "시점성 위반 0건"),
        ("일관성", "기업/종목/상장일 키가 단일 레코드로 정합", "매칭 격리 큐", "중복·모호 매칭률"),
        ("고유성", "rcept_no·원문 해시·스냅샷 ID 중복 금지", "중복 제거 및 감사 기록", "중복 레코드 0건"),
    ], [1300, 3350, 2850, 1860])
    add_callout(doc, "데이터 누수 금지", "종가 모델도 상장일 종가, 이후 체결, 사후 기사, 이후 수정 공시를 피처로 사용할 수 없다. 모든 피처는 예측이 생성된 시점보다 이른 수집·공시 시각을 증명해야 한다.")

    doc.add_heading("7. 보안, 접근제어, 감사", level=1)
    add_table(doc, ["영역", "필수 통제"], [
        ("인증", "관리자·운영자는 SSO 또는 MFA, 서비스 간에는 단기 토큰 또는 서비스 계정 사용"),
        ("권한", "RBAC: 일반 사용자=예측 조회, 운영자=상태 확인, 데이터 스튜어드=격리 검토, 모델 오너=후보 승인, 보안 오너=권한·비밀 관리"),
        ("비밀", "API 키·비밀번호는 Secret Manager에서 주입, 로그·오류·분석 도구에서 마스킹, 정기 회전"),
        ("암호화", "전송 TLS, 저장소 암호화, 백업 암호화, 키 접근 감사"),
        ("감사", "원문 접근, 데이터 변경, 권한 변경, 모델 승격, 예측 생성·삭제를 불변 감사 로그에 기록"),
        ("데이터 이용", "KRX·외부 원천별 계약·약관을 데이터 카탈로그에 연결하고 재배포 가능 범위를 API 응답 정책으로 강제"),
    ], [1900, 7460])

    doc.add_heading("8. 모델 거버넌스와 출시 통제", level=1)
    doc.add_paragraph("모델은 코드 파일이 아니라 버전이 있는 운영 자산으로 관리한다. 모델 레지스트리에는 학습 데이터 기간, 피처 집합, 결측 처리, 타깃 정의, 알고리즘·하이퍼파라미터, 검증 결과, 승인자, 배포·롤백 이력을 보관한다.")
    add_table(doc, ["게이트", "통과 조건"], [
        ("학습 데이터", "품질 상태가 PASS인 피처 스냅샷만 사용하고, 학습·검증·보정 구간의 시간 순서를 보존"),
        ("백테스트", "walk-forward 검증에서 현행 모델보다 MAE, 방향 정확도, 예측구간 포괄률을 종합 비교"),
        ("출시", "성능 개선, 누수 검사 통과, 모델 카드 작성, 모델 오너와 서비스 오너 승인"),
        ("운영", "실측 누적 오차·방향 정확도·커버리지·데이터 결측률을 감시"),
        ("롤백", "품질 게이트 실패, 성능 급락, 데이터 라이선스 이슈 시 이전 승인 모델 또는 예측 중지 상태로 즉시 전환"),
    ], [2000, 7360])

    doc.add_heading("9. 장애·변경·감사 대응", level=1)
    add_table(doc, ["사건", "즉시 조치", "후속 조치"], [
        ("DART/KRX 수집 실패", "해당 배치 FAIL, 이전 데이터를 새 데이터처럼 표시하지 않음, 알림", "재시도·원천 상태 확인·누락 범위 기록"),
        ("파서 변경", "새 버전은 과거 표본 회귀 테스트 후 제한 배포", "스키마·피처 영향과 이전 결과 차이 기록"),
        ("매칭 오류", "자동 예측/학습 대상에서 격리", "스튜어드 확인 후 수정 이력 보존"),
        ("비밀 노출 의심", "즉시 폐기·회전, 접근 토큰 무효화", "감사 범위 확인, 재발 방지 통제"),
        ("모델 이상", "신규 모델 중지 또는 롤백, 사용자에게 데이터 상태 표시", "원인 분석·재검증·승인 재수행"),
    ], [2000, 3600, 3760])

    doc.add_heading("10. 단계별 구현 로드맵", level=1)
    add_table(doc, ["단계", "완료 기준"], [
        ("0. 계약·접근 준비", "DART 키 발급, KRX 이용·재배포 범위 확인, Secret 저장소와 역할 정의"),
        ("1. 실제 이력 적재", "원문·정규화·정합·품질 요약 생성, 오류 격리 큐와 출처 메타데이터 확보"),
        ("2. 검증·모델", "실제 과거 데이터 walk-forward 성능 보고서, 모델 카드, 승격 기준 확정"),
        ("3. API·앱", "인증된 예측 조회, 입력 데이터 최신성·출처·불확실성 표시, 관리자 품질 화면"),
        ("4. 운영 자동화", "21:00 실측 작업, 알림, 재학습 정책, 백업·복구·감사 리허설"),
    ], [1900, 7460])
    add_callout(doc, "출시 전 최소 승인 체크", "데이터 이용권한 확인, 비밀관리 검증, 원문·피처·모델 계보 조회, 시점성 테스트, 실제 데이터 백테스트, 예측 면책·최신성 표시, 장애·롤백 리허설을 모두 통과해야 한다.")

    doc.add_heading("부록 A. 데이터셋 최소 스키마", level=1)
    add_table(doc, ["데이터셋", "핵심 키/필드", "소유자"], [
        ("raw_disclosure", "source, rcept_no, content_hash, collected_at, object_uri", "데이터 오너"),
        ("ipo_master", "ipo_id, corp_code, ticker, listing_date, match_confidence", "데이터 스튜어드"),
        ("feature_snapshot", "snapshot_id, ipo_id, as_of_at, feature_set_version, quality_status", "모델 오너"),
        ("prediction", "prediction_id, snapshot_id, model_version, open/close return, interval", "서비스 오너"),
        ("outcome_evaluation", "ipo_id, official_open, official_close, reconciled_at, error metrics", "모델 오너"),
        ("audit_event", "actor, action, resource, timestamp, result, request_id", "보안 오너"),
    ], [2200, 5300, 1860])

    doc.add_heading("부록 B. 참고 자료", level=1)
    add_source(doc, "SailPoint, 데이터 거버넌스 프레임워크: 소유권, 접근제어, 보존, 품질, 스튜어드십과 운영 모범사례 참고.", "https://www.sailpoint.com/ko/identity-library/creating-a-secure-data-governance-framework")
    add_source(doc, "Snowflake, 데이터 거버넌스 프레임워크: 전략, 책임성, 품질, 분류·메타데이터, 감사 증거의 운영 모델 참고.", "https://www.snowflake.com/ko/data-governance/frameworks/")
    add_source(doc, "OpenDART, 공시검색 API: 인증키와 공시 목록 수집 파라미터 참고.", "https://opendart.fss.or.kr/guide/detail.do?apiGrpCd=DS001&apiId=2019001")
    add_source(doc, "OpenDART, 공시서류 원본파일 API: document.xml의 ZIP 바이너리 응답과 인증키 요구사항 참고.", "https://opendart.fss.or.kr/guide/detail.do?apiGrpCd=DS001&apiId=2019003")
    add_source(doc, "KRX/KOSCOM, Market Data Usage Policies: 시장 데이터 외부 제공·재배포 관련 사전 검토 필요성 참고.", "https://data.krx.co.kr/inc/datasale/Market%20Data%20Usage%20Polices_ko.pdf?v=20230121_1")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
